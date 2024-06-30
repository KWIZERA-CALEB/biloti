from django.shortcuts import render, redirect
from .models import Student
from headteacher.models import ClassDetail, SchoolYear
from django.db.models import Q

from .forms import StudentAddForm
from django.http import HttpResponse

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor

from django.conf import settings

import os


from docx import Document
from io import BytesIO


# Create your views here.
def home(request):
    return render(request, 'teacher/home.html')

# def addStudent(request):

#     if request.method == "POST":
#         student_name = request.POST.get('student_name')
#         student_gender = request.POST.get('student_gender')
#         student_grade = request.POST.get('student_grade')

#         save_student = Student(
#             student_name=student_name,
#             student_gender=student_gender,
#             student_grade=student_grade
#         )

        

#         try:
#             save_student.save()
#             return redirect('classes')
#         except:
#             print('Failed to add student')
#             return redirect('add-student')

        
#     context = {}
#     return render(request, 'teacher/add_student.html', context)




def addStudent(request):
    form = StudentAddForm()
    if request.method == "POST":
        form = StudentAddForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('classes')
        
        
    context = {'form':form}
    return render(request, 'teacher/add_student.html', context)

def classes(request):
    # search_query = request.GET.get('search_query') if request.GET.get('search_query') != None else ''

    # search_query.replace(' ', '+')


    # # class_rooms = ClassDetail.objects.all()
    # class_rooms = ClassDetail.objects.filter(Q(class_name__icontains=search_query))

    # class_rooms_count = class_rooms.count()

    search_query = request.GET.get('search_query', '').strip()

    if search_query:
        search_words = search_query.split()
        query = Q()
        for word in search_words:
            query &= Q(class_name__icontains=word) # &= means query = query + Q(class_name__icontains=word)
        class_rooms = ClassDetail.objects.filter(query)
    else:
        class_rooms = ClassDetail.objects.all()

    class_rooms_count = class_rooms.count()

    

    context = {'class_rooms':class_rooms, 'class_rooms_count':class_rooms_count}
    return render(request, 'teacher/classes.html', context)

def classroom(request, pk):
    classes = ClassDetail.objects.get(id=pk)
    students = Student.objects.filter(student_grade=classes)
    students_class_rooms_count = students.count()

    
    context = {'classes':classes, 'students':students, 'students_class_rooms_count':students_class_rooms_count}
    return render(request, 'teacher/class-room.html', context)

def student(request, pk):
    student = Student.objects.get(id=pk)
    school_years = SchoolYear.objects.all()

    context = {'student':student,'school_years': school_years}
    return render(request, 'teacher/student-portal.html', context)

def generateDoc(request, pk):
    student = Student.objects.get(id=pk)
    
    
    if request.method == "POST":
        document = Document()


        # Adding Date paragraph
        date_paragraph = document.add_paragraph('2023/2024, Term 2')
        date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        date_run = date_paragraph.runs[0]
        date_run.font.size = Pt(16)
        date_run.bold = True
        date_run.font.name = 'Calibri'

        # add image
        image_path = os.path.join(settings.MEDIA_ROOT, student.student_avatar.name)

        student_avatar = document.add_picture(image_path, width=Inches(1))

        # Adding Students Names paragraph
        students_names_paragraph = document.add_paragraph()
        students_names_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        students_names_run = students_names_paragraph.add_run('NAMES:')
        students_names_run.font.size = Pt(16)
        students_names_run.bold = True
        students_names_run.font.name = 'Calibri'

        # Adding individual student names
        student_name_run = students_names_paragraph.add_run(f'{student.student_name}')
        student_name_run.font.size = Pt(16)
        student_name_run.font.underline = True
        student_name_run.bold = True
        student_name_run.font.color.rgb = RGBColor(10, 10, 200)
        student_name_run.font.name = 'Calibri'

        # Adding Grade paragraph
        student_grade_paragraph = document.add_paragraph(f'GRADE: {student.student_grade}')
        student_grade_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        student_grade_run = student_grade_paragraph.runs[0]
        student_grade_run.font.size = Pt(16)
        student_grade_run.bold = True
        student_grade_run.font.name = 'Calibri'

        # page break
        # document.add_page_break()


        teachers_comment_title = document.add_paragraph('Home room teacher’s comment')
        teachers_comment_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        teachers_comment_title_run = teachers_comment_title.runs[0]
        teachers_comment_title_run.font.size = Pt(16)
        teachers_comment_title_run.bold = True
        teachers_comment_title_run.font.name = 'Calibri'



        # create table
        table1 = document.add_table(rows=4, cols=4, style="Table Grid")

        # fill it 
        first_row = table1.rows[0].cells
        first_row[0].text = 'Activities'
        first_row[1].text = 'Usually'
        first_row[2].text = 'Sometimes'
        first_row[3].text = 'Rarely'

        second_row = table1.rows[1].cells
        second_row[0].text = 'Well-prepared and punctual'
        second_row[1].text = '✓'
        second_row[2].text = 'x'
        second_row[3].text = '✓'

        third_row = table1.rows[2].cells
        third_row[0].text = 'Completes assignments, meets deadlines'
        third_row[1].text = '✓'
        third_row[2].text = 'x'
        third_row[3].text = '✓'

        fourth_row = table1.rows[3].cells
        fourth_row[0].text = 'Follows class and school rules'
        fourth_row[1].text = '✓'
        fourth_row[2].text = 'x'
        fourth_row[3].text = '✓'


        # increase the width for column 1
        for row in table1.rows:
            row.cells[0].width = Inches(6)




        # table two
        table2 = document.add_table(rows=1, cols=4)
        row_one = table2.rows[0].cells
        row_one[0].text = 'Days: 6'
        row_one[1].text = 'Absent: 2'
        row_one[2].text = 'Late: 2'
        row_one[3].text = 'P.E.Non-Suit: NULL'




        comment_teachers_room = document.add_paragraph('Comment:  Anika has made great progress in her learning, now reading short sentences independently. She enjoys school activities with friends and am proud of her.Excited for next term.')
        comment_teachers_room.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        comment_teachers_room_run = comment_teachers_room.runs[0]
        comment_teachers_room_run.font.size = Pt(12)
        comment_teachers_room_run.font.name = 'Bookman Old'




        teachers_name = document.add_paragraph('Teacher Namara')
        teachers_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        teachers_name_run = teachers_name.runs[0]
        teachers_name_run.font.size = Pt(12)
        teachers_name_run.font.name = 'Bookman Old'

        school_name = document.add_paragraph('Two wings international school ')
        school_name.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        school_name_run = school_name.runs[0]
        school_name_run.font.size = Pt(14)
        school_name_run.bold = True
        school_name_run.font.name = 'Calibri'

        exam_details_title = document.add_paragraph('End of term 3 report: 2023-2024 Term 2')
        exam_details_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        exam_details_title_run = exam_details_title.runs[0]
        exam_details_title_run.font.size = Pt(12)
        exam_details_title_run.font.name = 'Bookman Old'

        # student_details = document.add_paragraph()
        # student_details.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # student_details_run = student_details.runs[1]
        # student_details_run.font.size = Pt(12)
        # student_details_run.font.name = 'Bookman Old'
        # student_details_add_names_title = student_details.add_run('Student:')
        # student_details_add_names = student_details.add_run('Anika Teta')

        library_title = document.add_paragraph('LIBRARY')
        library_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        library_title_run = library_title.runs[0]
        library_title_run.font.size = Pt(14)
        library_title_run.bold = True
        library_title_run.font.name = 'Calibri'

        # table3 = document.add_table(rows=2, cols=2)


        # Save the document
        # document.save('template2.docx')

        # use buffer to store doc in memory
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)


        
        response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={student.student_name} Report.docx'
        
        # Return the response to download report
        return response
    

    

